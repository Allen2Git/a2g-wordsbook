#!/usr/bin/env python
# -*- coding: utf-8 -*-
# @Time    : 2020, 04, 18, 20, 53
# @Author  : Allen Zhang
# @File    : ~/a2g-wordsbook/main/a2g-wordsbook.py
# @Description :
#
"""
```

usage: a2g-wordsbook.py [-h] [-i 输入文件名]
                        [--parse-type {txt,transcribe_json,pdf,word}]
                        [-b 熟词库文件名] [--base-type {txt,ddb}]
                        [-o 输出生词文件名]
                        [--out-type {pure_words,trans_words,youdao_xml,ddb}]

Wordsbook process the words file based on known words and output new words file

参数说明:
  -h, --help            显示帮助信息并退出
  -i PARSE_FILE_NAME, --parse-file PARSE_FILE_NAME
                        要处理的学习文件
  --parse-type {txt,transcribe_json,pdf,word}
                        支持的文件类型选择，可以是 txt-纯文本，transcribe-AWS Transcribe 翻译生成 JSON, pdf-PDF文件
  -b BASE_FILE, --base-file BASE_FILE
                        自己的熟词库文件，格式是每行一个单词
  --base-type {txt,ddb}
                        熟词库文件类型，可以是txt-纯文本，或者是 DynamoDB (开发中)
  -o WORDS_OUT_FILE, --words-out WORDS_OUT_FILE
                        输出的文件名
  --out-type {pure_words,trans_words,youdao_xml,ddb}
                        输出的文件类型，可以选择 pure_words-单纯每行单词的 txt 文本文件, trans_words-带有翻译的文本文件,
                        youdao_xml-可被导入的有道单词本XML文件, ddb 加入到 DynamoDB 词库当中
"""

# Default Constant Values:
import argparse  # main function use this module to be a commandline
from nltk.corpus import wordnet as wn
from os.path import dirname, join
import os
import json
import re
from io import BytesIO
from stardict import StarDict

# For parse PDF file:
from pdfminer.pdfinterp import PDFPageInterpreter, PDFResourceManager
from pdfminer.converter import TextConverter, PDFPageAggregator
from pdfminer.layout import LAParams
from pdfminer.pdfparser import PDFParser
from pdfminer.pdfdocument import PDFDocument
from pdfminer.pdfdevice import PDFDevice
from pdfminer.pdfpage import PDFPage

# For parse WORD file:
from docx import Document


# For create and parse XML
import xml.dom.minidom

# Define the word of threshold AWS Transcribe predicted confidence to choose
WORD_CONFIDENCE = 1
# new words file or txt file output to
NEW_WORDS_FILENAME = "./words/new_words.txt"
# base words lexicon it's mainly Allen's lexicon
BASE_WORDS_FILENAME = "./words/base_words.txt"
# Default or sample file to process
MEDIA_JSON_FILENAME = "./words/media-overview.json"
# Default Dict PATH
DICT_FILENAME = "./dict/star_sqldb.db"
# YOUDAO XML File Name
YOUDAO_XML_FILENAME = "./words/youdao_xml.xml"
# YOUDAO XML File Tags
YOUDAO_XML_TAGS = "a2g-wordsbook"

__version__ = '0.2'
__author__ = "Allen Zhang"


# class WordsParse used to parse input file including pure txt, pdf, transcribe json.
# The instance function will return the de-duplicated words of input file
# class paras as below:
# string(file_name) is input file name
# string(file_type) is one of txt, pdf, transcribe_json.
# return words = []

class WordsParse():
    def __init__(self,
                 parse_filename=MEDIA_JSON_FILENAME,
                 file_type="transcribe_json"
                 ):
        self.parse_filename = parse_filename
        self.file_type = file_type
        self.words = []
        self.current_dir = dirname(__file__)
        self.parse_filepath = join(self.current_dir, self.parse_filename)

        if self.file_type == "txt":
            self.txt_parse()
        elif self.file_type == "pdf":
            self.pdf_parse()
        elif self.file_type == "transcribe_json":
            self.transcribe_json_parse()
        elif self.file_type == "word":
            self.word_parse()
        else:
            # TODO throw exception
            pass

    # Check if the word or its morphy is a new word in base words
    def is_new_word(self, new_word=""):
        if (str.isalpha(new_word)):
            if wn.synsets(new_word, pos='v') == []:
                morphy_new_word = wn.morphy(new_word)
            else:
                morphy_new_word = wn.morphy(new_word, "v")
            if ((morphy_new_word) != None) and \
                    (morphy_new_word not in self.words):
                return morphy_new_word
        else:
            return False

    def pdf_parse(self):
        # parse the PDF file like paper to pure txt
        # Refer code from csdn, not dig too much about PDFMiner
        with open(self.parse_filepath, "rb") as parse_file:

            # 创建一个与文档相关的解释器
            parser = PDFParser(parse_file)

            # pdf文档的对象，与解释器连接起来
            doc = PDFDocument(parser=parser)
            parser.set_document(doc=doc)

            # 如果是加密pdf，则输入密码
            # doc._initialize_password()

            # 创建pdf资源管理器
            resource = PDFResourceManager()

            # 参数分析器
            laparam = LAParams()

            # 创建一个聚合器
            device = PDFPageAggregator(resource, laparams=laparam)

            # 创建pdf页面解释器
            interpreter = PDFPageInterpreter(resource, device)

            all_words = ""
            # 获取页面的集合
            for page in PDFPage.get_pages(parse_file):
                # 使用页面解释器来读取
                interpreter.process_page(page)

                # 使用聚合器来获取内容
                layout = device.get_result()
                for out in layout:
                    if hasattr(out, 'get_text'):
                        all_words += out.get_text()

            word_reg = re.compile(r'\w+')
            txt_file_words = word_reg.findall(all_words)
            for item in txt_file_words:
                new_word = str.lower(item)
                new_word = self.is_new_word(new_word)
                if (new_word and (new_word not in self.words)):
                    # add new word to new_words list
                    self.words.append(new_word)

    def transcribe_json_parse(self):
        with open(self.parse_filepath, "r") as parse_file:
            parse_file_json = json.loads(parse_file.read())
            for item in parse_file_json["results"]["items"]:
                new_word = str.lower(str(item["alternatives"][0]["content"]))
                word_confidence = float(item["alternatives"][0]["confidence"])
                if word_confidence >= WORD_CONFIDENCE:
                    # judge if it is a simple punctuation through word length.
                    new_word = self.is_new_word(new_word)
                    if (new_word and (new_word not in self.words)):
                        # add new word to new_words list
                        self.words.append(new_word)
                else:
                    continue

    def txt_parse(self):
        with open(self.parse_filepath, "r") as txt_file:
            word_reg = re.compile(r'\w+')
            txt_file_words = word_reg.findall(txt_file.read())
            for item in txt_file_words:
                new_word = str.lower(item)
                new_word = self.is_new_word(new_word)
                if (new_word and (new_word not in self.words)):
                    # add new word to new_words list
                    self.words.append(new_word)

    def word_parse(self):
        with open(self.parse_filepath, "rb") as word_file:
            source_stream = BytesIO(word_file.read())
            document = Document(source_stream)
            word_text = []
            for para in document.paragraphs:
                word_text.append(para.text)
            word_text = " ".join(word_text)
            word_reg = re.compile(r'\w+')

            source_stream.close()

            word_file_words = word_reg.findall(word_text)
            for item in word_file_words:
                new_word = str.lower(item)
                new_word = self.is_new_word(new_word)
                if (new_word and (new_word not in self.words)):
                    # add new word to new_words list
                    self.words.append(new_word)

# class Words Output used to output new words to different type like pure_words, \
# trans_words(word with translation), youdao_xml
# Words output direct to a file
class WordsOutput():
    def __init__(self,
                 new_words=[],
                 output_filename=NEW_WORDS_FILENAME,
                 output_type="pure_words"

                 ):
        self.new_words = new_words
        self.output_type = output_type
        self.output_filename = output_filename
        self.current_dir = dirname(__file__)
        self.parse_file_path = join(self.current_dir, self.output_filename)

        if output_type == "pure_words":
            self.pure_words_output()
        elif output_type == "trans_words":
            self.trans_words_output()
        elif output_type == "youdao_xml":
            self.youdao_xml_output()
        elif output_type == "ddb":
            self.ddb_output()
        else:  # TODO throw exception for unknown param
            pass

    def pure_words_output(self):
        new_words_str = ""
        with open(self.parse_file_path, "w") as output_words_file:
            for word in self.new_words:
                new_words_str += word + "\n"
            output_words_file.write(new_words_str)

    def trans_words_output(self):
        new_words_str = ""
        dict_file_path = join(self.current_dir, DICT_FILENAME)
        # Use Stardict refer https://github.com/skywind3000/ECDICT
        dict = StarDict(dict_file_path)
        with open(self.parse_file_path, "w") as output_words_file:
            for word in self.new_words:
                word_output = dict.query(word)
                # Output format as below
                out_put = "{} {} \n {} \n".format(
                    word_output['word'], word_output['phonetic'], word_output['translation']
                )
                new_words_str += out_put + "\n"
            output_words_file.write(new_words_str)

    def youdao_xml_output(self):
        ''' example as below:
        <item>    
        <word>phonetic</word>
        <trans><![CDATA[adj. 语音的，语音学的；音形一致的；发音有细微区别的]]></trans>
        <phonetic><![CDATA[[fə'netɪk]]]></phonetic>
        <tags>未分组</tags>
        <progress>1</progress>
        </item>
        '''
        # Open Dict
        dict_file_path = join(self.current_dir, DICT_FILENAME)
        # Use Stardict refer https://github.com/skywind3000/ECDICT
        dict = StarDict(dict_file_path)

        youdao_xml = xml.dom.minidom.Document()
        root = youdao_xml.createElement('wordbook')
        youdao_xml.appendChild(root)
        for word in self.new_words:
            word_query = dict.query(word)
            node_item = youdao_xml.createElement('item')
            
            node_item_word = youdao_xml.createElement("word")
            node_item_word.appendChild(youdao_xml.createTextNode(word))
            node_item.appendChild(node_item_word)

            node_item_trans = youdao_xml.createElement("trans")
            node_item_trans.appendChild(youdao_xml.createTextNode(word_query['translation']))
            node_item.appendChild(node_item_trans)

            node_item_phonetic = youdao_xml.createElement("phonetic")
            node_item_phonetic.appendChild(youdao_xml.createTextNode(word_query['phonetic']))
            node_item_phonetic.appendChild(youdao_xml.createCDATASection(word_query['phonetic']))
            node_item.appendChild(node_item_phonetic)

            node_item_tags = youdao_xml.createElement("tags")
            node_item_tags.appendChild(youdao_xml.createTextNode(YOUDAO_XML_TAGS))
            node_item.appendChild(node_item_tags)

            node_item_progress = youdao_xml.createElement("progress")
            node_item_progress.appendChild(youdao_xml.createTextNode("1"))
            node_item.appendChild(node_item_progress)

            root.appendChild(node_item)
            
        with open(self.parse_file_path, "w") as output_words_file:
            youdao_xml.writexml(output_words_file, indent='\t', addindent='\t', newl='\n', encoding="utf-8")

    def ddb_output(self):
        pass

# Base Words is Allen's lexicon contained words
# self.base_words is the most often output word list.


class BaseWords():
    def __init__(self,
                 base_words_file=BASE_WORDS_FILENAME,
                 base_words_type="txt"
                 ):
        self.base_words = []
        self.base_words_filename = base_words_file
        self.base_words_type = base_words_type

        self.current_dir = dirname(__file__)
        self.base_words_file_path = join(
            self.current_dir, self.base_words_filename)

        if base_words_type == "txt":
            self.get_from_txt()
        elif base_words_type == "ddb":
            self.get_from_ddb()
        else:  # TODO throw new exception
            pass

    def get_from_txt(self):
        with open(self.base_words_file_path, "r") as base_words_file:
            base_words_lines = base_words_file.read().split()
            word = ""
            for word in base_words_lines:
                self.base_words.append(word)

    def get_from_ddb(self):
        pass

    def add_new_words(self):
        pass

    def add_recited_words(self):
        pass


# Process file directly using command line. Format as below:
# wordsbook --parse-file <filename> --parse-type <txt, transcribe_json or pdf> --word-out <filename>
# 	--out-type <pure_words, trans_words, yudao_xml, ddb>

def wordsbook_cmd():
    cmd_parse = argparse.ArgumentParser(
        description="Wordsbook process the words file and output new words file"
    )
    cmd_parse.add_argument('-i', '--parse-file', type=str, default='input_file.txt',
                           action='store', dest='parse_file_name',
                           help='File name to parse'
                           )
    cmd_parse.add_argument('--parse-type', type=str,
                           action='store', dest='parse_type',
                           default="txt", choices=['txt', 'transcribe_json', 'pdf', 'word'],
                           help='Input file parse type, can be txt, transcribe_json, pdf,word'
                           )
    cmd_parse.add_argument('-b', '--base-file', type=str,default='base_words.txt',
                           action='store', dest='base_file',
                           help='Lexicon file name'
                           )
    cmd_parse.add_argument('--base-type', type=str,
                           action='store', dest='base_type',
                           default="txt", choices=['txt',  'ddb'],
                           help='Base file can be txt or from ddb'
                           )
    cmd_parse.add_argument('-o', '--words-out', type=str, default='new_words.txt',
                           action='store', dest='words_out_file',
                           help='Output file name'
                           )
    cmd_parse.add_argument('--out-type', action='store', type=str,
                           dest='out_type', default="pure_words",
                           choices=['pure_words', 'trans_words',
                                    'youdao_xml', 'ddb'],
                           help='Output file name can be pure_words, trans_words, youdao_xml, ddb'
                           )
    cmd_parse = cmd_parse.parse_args()

    print('{} is parsing...\n'.format(cmd_parse.parse_file_name))
    words_parse = WordsParse(parse_filename=cmd_parse.parse_file_name,
                             file_type=cmd_parse.parse_type
                             )
    words_parse = words_parse.words

    print('{} as base words file is used\n'.format(cmd_parse.base_file))
    base_words = BaseWords(cmd_parse.base_file, cmd_parse.base_type)
    base_words = base_words.base_words
    new_words = []

    for word in words_parse:
        if word not in base_words:
            new_words.append(word)
        else:
            continue

    WordsOutput(new_words, cmd_parse.words_out_file, cmd_parse.out_type)
    print('output file {} is generated\n'.format(cmd_parse.words_out_file))

if __name__ == '__main__':
    wordsbook_cmd()
